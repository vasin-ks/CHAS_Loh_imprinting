#  Input example:
# 14 columns from the program ", "1:PatientName 2:Type 3:Chromosome 4:Min 5:Max 6:Size 7:M.Count 8:M.M.Dis 9:Cyto.Start\
# 10:Cyto.End 11:Gene.Count 12:Omim.G.Count 13:Genes 14:Omim.Genes")
#  Name_off_pattient LOH	4	126303972	128296809	1992.838	437	4570	q28.1	q28.2	8	6	INTU, SLC25A31,\
#  HSPA4L, PLK4, MFSD8, ABHD18, LARP1B, PGRMC2	INTU (610621), SLC25A31 (610796), HSPA4L (619077), PLK4 (605031), MFSD8\
#  (611124), PGRMC2 (607735)

#install python-docx in PyCharm
from docx import Document
from docx.shared import RGBColor
document = Document()
red = RGBColor(255, 0, 0)

# www.geneimprint.com from this site you can add new imprinting genes or introduce additional genes from another
geneimprint = ["DIRAS3", "ARHI", "NOEY2", "RNU5D-1", "U5DL", "U5DS", "RNU5D", "TP73", "P73", "LRRTM1", "ZDBF2", "GPR1", "NAP1L5", "DRLM", "ERAP2", "LRAP", "L-RAP", "RHOBTB3", "VTRNA2-1", "CBL3", "CBL-3", "hvg-5", "nc886", "MIR886", "VTRNA2", "MIRN886", "svtRNA2-1a", "hsa-mir-886", "ADTRP", "AIG1L", "C6orf105", "dJ413H6.1", "FAM50B", "X5L", "D6S2654E", "PXDC1", "C6orf145", "LIN28B", "CSDD2", "AIM1", "ST4", "CRYBG1", "PLAGL1", "ZAC", "LOT1", "ZAC1", "MGC126275", "MGC126276", "DKFZp781P1017", "HYMAI", "NCRNA00020", "SLC22A2*", "OCT2", "MGC32628", "SLC22A3*", "EMT", "EMTH", "OCT3", "GRB10", "RSS", "IRBP", "MEG1", "GRB-IR", "Grb-10", "KIAA0207", "DDC", "AADC", "HECW1", "NEDL1", "MAGI2", "AIP1", "AIP-1", "ARIP1", "SSCAM", "MAGI-2", "ACVRIP1", "PEG10", "EDR", "HB-1", "Mar2", "MEF3L", "Mart2", "RGAG3", "SGCE", "ESG", "DYT11", "PPP1R9A", "NRB1", "NRBI", "FLJ20068", "KIAA1222", "Neurabin-I", "TFPI2", "PP5", "REF1", "TFPI-2", "FLJ21164", "DLX5", "CCDC71L", "C7orf74", "COPG2IT1", "CIT1", "COPG2AS", "FLJ41646", "NCRNA00170", "DKFZP761N09121", "CPA4", "CPA3", "MEST", "PEG1", "MGC8703", "MGC111102", "DKFZp686L18234", "MESTIT1", "MEST-IT", "PEG1-AS", "MEST-AS1", "MEST-IT1", "NCRNA00040", "KLF14", "BTEB5", "DLGAP2", "DAP2", "SAPAP2", "KCNK9", "KT3.2", "TASK3", "K2p9.1", "TASK-3", "MGC138268", "MGC138270", "ZFAT-AS1", "ZFATAS", "ZFAT-AS", "SAS-ZFAT", "NCRNA00070", "ZFAT", "AITD3", "ZFAT1", "ZNF406", "PEG13", "GLIS3", "ZNF515", "INPP5F", "V2", "SAC2", "hSAC2", "MSTP007", "MSTPO47", "FLJ13081", "KIAA0966", "MGC59773", "MGC131851", "WT1", "GUD", "AWT1", "WAGR", "WT33", "NPHS4", "WIT-2", "EWS-WT1", "WT1-AS", "WIT1", "WIT-1", "WT1AS", "WT1-AS1", "KCNQ1OT1", "LIT1", "KvDMR1", "KCNQ10T1", "KvLQT1-AS", "OSBPL5", "ORP5", "OBPH1", "FLJ42929", "KCNQ1DN", "BWRT", "HSA404617", "SLC22A18", "HET", "ITM", "BWR1A", "IMPT1", "TSSC5", "ORCTL2", "BWSCR1A", "SLC22A1L", "p45-BWR1A", "DKFZp667A184", "IGF2", "INSIGF", "pp9974", "C11orf43", "FLJ22066", "FLJ44734", "IGF2AS", "PEG8", "MGC168198", "PHLDA2", "IPL", "BRW1C", "BWR1C", "HLDA2", "TSSC3", "CDKN1C", "BWS", "WBS", "p57", "BWCR", "KIP2", "KCNQ1", "LQT", "RWS", "WRS", "LQT1", "SQT2", "ATFB1", "ATFB3", "JLNS1", "KCNA8", "KCNA9", "Kv1.9", "Kv7.1", "KVLQT1", "FLJ26167", "H19", "ASM", "BWS", "ASM1", "MGC4485", "PRO2605", "D11S813E", "INS", "ILPR", "IRDN", "ANO1", "DOG1", "TAOS2", "ORAOV2", "TMEM16A", "ZC3H12C", "MCPIP3", "NTM", "HNT", "NTRI", "IGLON2", "ST8SIA1", "GD3S", "SIAT8", "SIAT8A", "SIAT8-A", "ST8SiaI", "RBP5", "CRBP3", "CRBPIII", "CRBP-III",  "HNF1A", "HNF1", "LFB1", "TCF1", "HNF4A", "MODY3", "TCF-1", "HNF-1A", "IDDM20", "ATP5F1EP2", "ATP5EP2", "RB1", "RB", "pRb", "OSRC", "pp110", "p105-Rb", "ESR2", "Erb", "ESRB", "ODG8", "ESTRB", "NR3A2", "ER-BETA", "ESR-BETA", "SMOC1", "OAS", "MEG3", "GTL2", "FP504", "prebp1", "PRO0518", "PRO2160", "FLJ31163", "FLJ42589", "DIO3", "D3", "5DIII", "TXDI3", "DIOIII", "DLK1", "DLK", "FA1", "ZOG", "pG2", "DLK-1", "PREF1", "Delta1", "Pref-1", "MEG8", "Bsr", "Irm", "Rian", "SNHG23", "SNHG24", "lnc-MGC", "LINC00024", "NCRNA00024", "DIO3OS", "DIO3-OS", "DIO3-AS1", "C14orf134", "NCRNA00041", "SNORD113-1", "14q(I-1)", "SNORD114-1", "14q(II-1)", "RTL1", "MART1", "PEG11", "LOC388015", "MAGEL2", "nM15", "NDNL1", "UBE3A", "AS", "ANCR", "E6-AP", "HPVE6A", "EPVE6AP", "FLJ26981", "MKRN3", "D15S9", "RNF63", "ZFP127", "ZNF127", "MGC88288", "SNORD116", "PET1", "PWCR1", "HBII-85", "SNORD115@", "HBII-52", "NPAP1", "C15orf2", "SNORD109B", "HBII-438B", "HBII-438B", "C/D", "box", "snoRNA", "SNORD109A", "HBII-438A", "PWAR6", "HBT8", "PAR-6", "ATP10A", "ATPVA", "ATPVC", "ATP10C", "KIAA0566", "SNORD108", "HBII-437", "HBII-437", "C/D", "box", "snoRNA", "SNORD115-48", "HBII-52-48", "SNORD107", "HBII-436", "HBII-436", "C/D", "box", "snoRNA", "PWCR1", "PET1", "non-coding", "RNA", "in", "the", "Prader-Willi", "critical", "region", "SNRPN", "SMN", "PWCR", "SM-D", "RT-LI", "HCERN3", "SNRNP-N", "FLJ33569", "FLJ36996", "FLJ39265", "MGC29886", "SNURF-SNRPN", "DKFZp762N022", "DKFZp686C0927", "DKFZp761I1912", "DKFZp686M12165", "NDN", "HsT16328", "SNORD64", "HBII-13", "HBII-13", "snoRNA", "SNURF", "RASGRF1", "GNRP", "GRF1", "CDC25", "GRF55", "CDC25L", "H-GRF55", "PP13187", "IRAIN", "IGF1R-AS", "NAA60", "HAT4", "NAT15", "ZNF597", "CMTM1", "CKLFH", "CKLFH1", "CKLFSF1", "ZFP90", "FIK", "NK10", "ZNF756", "zfp-90", "TP53", "P53", "BCC7", "LFS1", "TRP53", "ZNF396", "ZSCAN14", "TCEB3C", "HsT829", "TCEB3L2", "Elongin", "A3", "DNMT1", "AIM", "DNMT", "MCMT", "CXXC9", "HSN1E", "ADCADN", "MIMT1", "MIM1", "LINC00067", "NCRNA00067", "ZIM2", "ZNF656", "PEG3", "PW1", "ZNF904", "ZSCAN24", "NLRP2", "NBS1", "PAN1", "NALP2", "PYPAF2", "CLR19.9", "MIR371A", "C19MC", "MIR371", "MIRN371", "hsa-mir-371", "hsa-mir-371a", "PEG3-AS1", "APEG3", "PEG3AS", "PEG3-AS", "NCRNA00155", "PSIMCT-1", "NNAT", "Peg5", "BLCAP", "BC10", "MCTS2", "GDAP1L1", "dJ881L22.1", "dJ995J12.1.1",  "SGK2", "H-SGK2", "dJ138B7.2", "GNAS", "AHO", "GSA", "GSP", "POH", "GPSA", "NESP", "GNAS1", "PHP1A", "PHP1B", "C20orf45", "MGC33735", "dJ309F20.1.1", "dJ806M20.3.3", "L3MBTL1", "L3MBTL", "ZC2HC3", "H-L(3)MBT", "dJ138B7.3", "GNASAS", "SANG", "NESPAS", "GNAS1AS", "NCRNA00075", "SANG", "SANG", "Nespas", "MIR296", "MIRN296", "miRNA296", "MIR298", "MIRN298", "hsa-mir-298", "DSCAM", "CHD2", "CHD2-42", "CHD2-52", "DGCR6L", "DGCR6"]

# Enter how much the "CHAS" program gives you "LOH" for the patient
print("Enter the total number of LOH")
import re
m = int(input())
print("Copy", m, "LOH from CHAS")
print("14 columns from the program, ", "one line example:", "1:PatientName 2:Type 3:Chromosome 4:Min 5:Max 6:Size \
7:M.Count 8:M.M.Dis 9:Cyto.Start 10:Cyto.End 11:Gene.Count 12:Omim.G.Count 13:Genes 14:Omim.Genes") #


for i in range(m):
    i = input().split()
    Loh = i
    VsegoGenov = i[10]
    VsegoGenov = int(VsegoGenov)
    geni = i[12:VsegoGenov + 12]
    geni=[re.sub(',*','',i) for i in geni]
    size = i[5]
    size = size.replace(".", "")
    size = int(size)
    ngen = i[10]
    b = len(geni)
    imp = []
    N_imp_gen = 0
    razloh = 0
    cit = 12
    if i[8] == i[9]:
        cit = i[8]
    if i[8] != i[9]:
        cit = i[8],i[9]
    cit = "".join(cit)
    cor = (i[2],cit, "(", i[3], "_", i[4], ")",)
    cor = "".join(cor)
    for i in range(b):
        if geni[i] in geneimprint:
            imp += [geni[i]]
            N_imp_gen += 1
        if size > 2500000: # LOH size, above which we will consider the LOH important for additional analysis
            razloh += 1

    if N_imp_gen > 0 and razloh > 0:
        Loh = ' '.join(Loh)
        Loh = Loh.replace('(', '[OMIM:')
        Loh = Loh.replace(')', ']')
        p = document.add_paragraph(' ')
        p.add_run('Imprinting Genes + Size:  ').font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(','.join(imp)).font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(' and ')
        p.add_run(str(size)).font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(' ' + Loh)
        p.add_run(' ' + cor).font.color.rgb = RGBColor(255, 0, 0)
    elif N_imp_gen > 0 and razloh == 0:
        Loh = ' '.join(Loh)
        Loh = Loh.replace('(', '[OMIM:')
        Loh = Loh.replace(')', ']')
        p = document.add_paragraph(' ')
        p.add_run("Imprinting Genes: ").font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(','.join(imp)).font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(' ' + Loh)
        p.add_run(' ' + cor).font.color.rgb = RGBColor(255, 0, 0)
    elif N_imp_gen == 0 and razloh > 0:
        Loh = ' '.join(Loh)
        Loh = Loh.replace('(', '[OMIM:')
        Loh = Loh.replace(')', ']')
        p = document.add_paragraph(' ')
        p.add_run('Size: ').font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(str(size)).font.color.rgb = RGBColor(255, 0, 0)
        p.add_run(' ' + Loh)
        p.add_run(' ' + cor).font.color.rgb = RGBColor(255, 0, 0)

document.save(r"C:\Users\Vasin\Desktop\imprinting_analysis .doc")  #write your path to save the analysis results
